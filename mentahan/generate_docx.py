from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys

def main():
    doc = Document()
    
    # Title
    title = doc.add_heading('PROPOSAL PROYEK AKHIR ARTIFICIAL INTELLIGENCE\n', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    runner = subtitle.add_run('SMART DOCUMENT CONVERTER BERBASIS AI\n(Pemahaman Struktur dan Rekonstruksi Layout)\n')
    runner.bold = True
    runner.font.size = Pt(14)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    team_heading = doc.add_paragraph()
    team_run = team_heading.add_run('Disusun Oleh Kelompok:')
    team_run.bold = True
    
    doc.add_paragraph('1. Ramadhan Al Fitra', style='List Number')
    doc.add_paragraph('2. Taufiq', style='List Number')
    doc.add_paragraph('3. Zikri', style='List Number')
    
    doc.add_page_break()

    sections = [
        ("1. Analisis Permasalahan", [
            "Saat ini, berbagai platform konversi dokumen konvensional (seperti iLovePDF, SmallPDF, dll) melakukan konversi file secara langsung berdasarkan translasi format asal ke format tujuan tanpa benar-benar 'memahami' isi dan konteks dokumen tersebut.",
            "Permasalahan utama yang sering terjadi meliputi:",
            "- Kerusakan layout, margin, dan perataan teks (alignment).",
            "- Pergeseran posisi tabel, perubahan ukuran font secara tidak wajar, serta hilangnya elemen krusial seperti header/footer.",
            "- Kegagalan mengekstrak teks pada dokumen hasil scan (terutama file PDF hasil pindaian mesin scanner/foto).",
            "- Tidak adanya rekonstruksi struktur hierarki dokumen (seperti Heading, Paragraf, dan List) sehingga dokumen hasil konversi menjadi sangat sulit untuk diedit kembali secara rapi di Microsoft Word atau software serupa."
        ]),
        ("2. Solusi AI", [
            "Proyek ini mengusulkan pembangunan 'Smart Document Converter', yakni sebuah sistem berbasis Artificial Intelligence (Kecerdasan Buatan).",
            "Dalam proyek ini, AI tidak hanya digunakan sebagai fitur tambahan semata, melainkan menjadi mesin utama (core engine) dalam keseluruhan siklus pemrosesan file. Solusi AI yang ditawarkan meliputi:",
            "- AI Document Analysis & Classification: AI akan membaca dan mengelompokkan jenis dokumen, serta secara cerdas mendeteksi elemen struktural (paragraf, tabel, gambar, numbering, hyperlink).",
            "- Layout Detection & Recovery: Memanfaatkan algoritma Computer Vision untuk mengenali penempatan letak (layout) elemen asli, sehingga format tidak berantakan pasca-konversi.",
            "- Smart OCR (Optical Character Recognition): Memanfaatkan model Deep Learning untuk mengonversi gambar atau dokumen berbasis pixel menjadi teks terstruktur yang dapat diedit secara bebas.",
            "- Smart Formatting Reconstructor: Merekonstruksi ulang struktur tata letak (seperti penyesuaian otomatis terhadap ukuran font dan pemulihan tabel yang rusak) agar sedekat mungkin dengan dokumen aslinya."
        ]),
        ("3. Analisis Kebutuhan", [
            "A. Kebutuhan Perangkat Keras (Hardware)",
            "- Server Backend dengan dukungan komputasi GPU yang memadai untuk melakukan inferensi model AI Computer Vision dan NLP (misal: Nvidia T4 / A100 di lingkungan cloud).",
            "- Minimal RAM 16 GB - 32 GB untuk memproses dokumen berjumlah ratusan halaman secara asinkron tanpa terjadinya kegagalan memori (Out Of Memory).",
            "B. Kebutuhan Perangkat Lunak (Software)",
            "- Sistem Operasi: Distro Linux (Ubuntu Server) direkomendasikan untuk stabilitas model AI.",
            "- Bahasa Pemrograman Utama: Python (sebagai Engine AI dan Backend API), JavaScript/HTML/CSS (untuk membangun antarmuka web interaktif Frontend).",
            "- Framework Pendukung: Laravel/Node.js untuk orkestrasi web request, dan Python Flask/FastAPI untuk microservice AI.",
            "- Pustaka AI & Machine Learning: HuggingFace Transformers, LangChain, OpenCV, Tesseract OCR / PaddleOCR.",
            "- Pustaka Konversi Dokumen Dasar: python-docx, PyMuPDF, Pandoc, LibreOffice (unoconv/headless).",
            "C. Kebutuhan Data",
            "- Dataset sampel dokumen berformat DOCX, XLSX, PDF, PPTX dengan berbagai variasi layout yang kompleks guna pengujian model (QA & Fine-tuning)."
        ]),
        ("4. Functional Requirement", [
            "1. Sistem harus menyediakan fungsionalitas untuk menerima unggahan file pengguna dalam format DOCX, XLSX, PPTX, dan PDF.",
            "2. Sistem harus mengizinkan pengguna memilih jenis target konversi di antara kombinasi keempat format tersebut secara dinamis (Contoh: Word ke PDF, PDF ke Excel, Excel ke PPT, dll).",
            "3. Sistem AI Engine harus menganalisis struktur hierarki isi dokumen (Heading, Tabel, List, Gambar) sebelum eksekusi konversi.",
            "4. Sistem harus secara otomatis mengaktifkan AI OCR jika mendeteksi masukan dokumen berupa gambar atau PDF berbasis piksel murni.",
            "5. Sistem harus memuat modul Smart Formatting yang berfungsi memverifikasi dan memperbaiki format (font, margin) paska proses translasi data.",
            "6. Sistem harus menampilkan layar pratinjau (preview) ringkasan dokumen atau analisis AI kepada pengguna.",
            "7. Sistem harus menyediakan tautan unduh (download) dan kemampuan merekam riwayat konversi tiap pengguna secara aman."
        ]),
        ("5. Non Functional Requirement", [
            "1. Performance: Waktu rata-rata proses unggah, analisis AI, hingga siap unduh harus berada di bawah batas waktu 45 detik untuk dokumen bervolume menengah (di bawah 30 halaman).",
            "2. Usability: Antarmuka sistem harus memiliki prinsip desain yang clean, intuitif, dan tidak membutuhkan panduan teknis yang rumit agar semua kalangan pengguna dapat memahaminya.",
            "3. Security & Privacy: Sistem harus dilengkapi dengan fitur penghapusan otomatis (auto-deletion) dari server secara periodik (maksimal 24 jam) untuk menggaransi kerahasiaan data pengguna.",
            "4. Scalability: Pemisahan arsitektur antara layanan antarmuka (Web Backend) dengan layanan komputasi (AI API Server) untuk memfasilitasi horizontal scaling manakala terjadi lonjakan pengguna secara masif."
        ]),
        ("6. Use Case Diagram", [
            "Aktor yang terlibat dalam sistem adalah Pengguna (User) dan Modul AI (AI Engine).",
            "Daftar Usecase Sistem Meliputi:",
            "- Mengakses Halaman Utama Converter.",
            "- Memilih Jenis Konversi dan Mengunggah File (User).",
            "- Menganalisis Struktur File (AI Engine).",
            "- Memproses Eksekusi OCR dan Smart Formatting (AI Engine).",
            "- Melihat Ringkasan Pratinjau (User).",
            "- Mengunduh Dokumen Final (User).",
            "- Melihat Log/Riwayat Konversi di Halaman Riwayat (User)."
        ]),
        ("7. Use Case Description", [
            "Use Case: Mengunggah File Dokumen",
            "- Aktor: Pengguna",
            "- Deskripsi Singkat: Proses di mana pengguna menentukan parameter tipe konversi lalu mengunggah file dari penyimpanan lokal.",
            "- Pre-condition: Pengguna berada di halaman Converter dan layanan backend berstatus online.",
            "- Post-condition: File terunggah ke temporary storage, dan status konversi tercatat 'Pending' di database.",
            "",
            "Use Case: Menjalankan Smart Formatting & Konversi",
            "- Aktor: AI Engine",
            "- Deskripsi Singkat: Sistem kecerdasan buatan membaca metadata file, melakukan perbaikan layout dan translasi ke bentuk dokumen baru.",
            "- Pre-condition: File pengguna telah tervalidasi dengan aman.",
            "- Post-condition: Menghasilkan file output yang secara visual mereplikasi file asal pada format ekstensi baru."
        ]),
        ("8. Activity Diagram", [
            "Runtutan Aktivitas Utama (End-to-End Workflow):",
            "1. User mengunjungi alamat Web Smart Converter.",
            "2. User memilih opsi tujuan konversi (contoh: PPT to Word).",
            "3. User memilih dan melakukan tindakan (action) unggah file lokal.",
            "4. Web Frontend mengalihkan tampilan ke loading bar dan mengirim data ke Web Backend.",
            "5. Web Backend meneruskan muatan (payload) menuju AI Engine untuk dieksekusi.",
            "6. AI Engine melakukan siklus: AI Analysis -> Layout Detection -> Konversi -> Smart Formatting.",
            "7. AI Engine menyelesaikan siklus, mengembalikan file baru ke Web Backend.",
            "8. Web Backend menyimpan log riwayat dan merespon Frontend dengan data Pratinjau (Preview).",
            "9. User disajikan Halaman Preview dan mengklik tautan Download.",
            "10. Sistem mencatat keberhasilan eksekusi tersebut (Finish)."
        ]),
        ("9. Sequence Diagram", [
            "Komponen Terkait: User, Web UI (Frontend), Web API (Backend), AI Microservice, Database Server.",
            "Rangkaian Interaksi:",
            "- User -> Web UI : Pilih File dan Eksekusi Upload",
            "- Web UI -> Web API : Request POST /api/upload-and-convert",
            "- Web API -> Database Server : Insert Data Histori Konversi (Status: Processing)",
            "- Web API -> AI Microservice : Send Document Object Stream",
            "- AI Microservice -> AI Microservice : Run Document Classification & Structure Parsing",
            "- AI Microservice -> AI Microservice : Run Model Inference (OCR/Table Detection if needed)",
            "- AI Microservice -> Web API : Return Formatted Output File",
            "- Web API -> Database Server : Update Data Histori (Status: Success/Completed)",
            "- Web API -> Web UI : Return File Link & AI Summary Data",
            "- Web UI -> User : Tampilkan Tombol Unduh dan Rangkuman"
        ]),
        ("10. Class Diagram", [
            "Representasi logikal melalui Class:",
            "Class [User]",
            "- Attributes: id_user (String), session_id (String), created_at (Date)",
            "- Methods: viewHistory()",
            "",
            "Class [Dokumen_Proses]",
            "- Attributes: id_dokumen (String), nama_asli (String), format_asli (String), size (Double), jenis_konversi (String)",
            "- Methods: validateFile(), compressFile()",
            "",
            "Class [Conversion_Log]",
            "- Attributes: id_log (Int), id_user (String), id_dokumen (String), status (String), ai_summary (Text), timestamp (Date)",
            "- Methods: recordLog(), updateStatus()",
            "",
            "Class [AI_Processor]",
            "- Attributes: model_version (String), is_gpu_active (Boolean)",
            "- Methods: executeLayoutRecovery(), runPaddleOCR(), generateSmartFormatting()"
        ]),
        ("11. ERD (Entity Relationship Diagram)", [
            "Entitas Dasar: USER, FILE_DOKUMEN, RIWAYAT_KONVERSI",
            "Relasi Kardinalitas:",
            "- Entitas USER (1) mencatat (M) transaksi RIWAYAT_KONVERSI.",
            "- Entitas RIWAYAT_KONVERSI (1) terkait dengan (1) FILE_DOKUMEN.",
            "",
            "Spesifikasi Atribut Utama:",
            "- USER: user_id (Primary Key), user_ip, session_token.",
            "- RIWAYAT_KONVERSI: log_id (Primary Key), user_id (Foreign Key), tipe_konversi, waktu_eksekusi, status_berhasil.",
            "- FILE_DOKUMEN: file_id (Primary Key), log_id (Foreign Key), nama_file_asli, nama_file_hasil, path_output, ukuran_file_kb."
        ]),
        ("12. Flowchart Sistem", [
            "1. START",
            "2. Tampilkan Halaman Beranda / Home",
            "3. Navigasi ke Halaman Converter",
            "4. Drop/Upload File & Tentukan Parameter Konversi",
            "5. Pengecekan Validitas Ekstensi dan Batasan Ukuran Maksimal",
            "6. Jika Invalid: Menampilkan Peringatan UI, kembali ke langkah 4.",
            "7. Jika Valid: Meneruskan stream data memanggil API Backend -> AI Engine Server.",
            "8. Tampilkan progress indikator. (Tunggu Proses AI Selesai).",
            "9. Beralih ke Halaman Preview (Menyajikan status 'Success' dan ringkasan ekstrasi struktur dokumen).",
            "10. Eksekusi proses Download file konversi akhir ke perangkat klien.",
            "11. Mencatat aktivitas ini dalam Session User di Halaman Riwayat Konversi.",
            "12. END"
        ]),
        ("13. Flowchart AI", [
            "1. START AI PROCESSING WORKER",
            "2. Menerima muatan dokumen biner",
            "3. Document Preprocessing: Mengurai struktur dasar metadata dan pengenalan parser.",
            "4. AI Document Analysis: Mendeteksi apakah layout bersifat kompleks (kolom ganda, gambar dominan).",
            "5. Logika Percabangan: Apakah isi file dominan berbasis piksel scan/gambar?",
            "   - Jika YA: Masuk jalur model OCR (Vision based pipeline - PaddleOCR/Tesseract).",
            "   - Jika TIDAK: Lanjut ke ekstraksi hirarki teks native.",
            "6. AI Layout Detection & Classification: Membedakan objek Header, Footer, Heading Title, Paragraf Dasar, List Numbering, dan Tabel Struktur matriks.",
            "7. Smart Translation ke objek internal virtual (In-memory Object Mapping).",
            "8. Rekonstruksi Format dan Rendering File Target: Menyusun tata letak kembali agar ukuran tabel tidak terpotong dan margin aman.",
            "9. AI Validation & Exporting: Menyimpan data byte file akhir (.docx / .pdf).",
            "10. Outputkan ke HTTP response.",
            "11. END AI PROCESSING WORKER"
        ]),
        ("14. Arsitektur Sistem", [
            "Model arsitektur berbasis Microservices (Layanan Terdistribusi):",
            "1. Presentation Layer (Client Browser): Dibangun menggunakan HTML, CSS modern (Tailwind atau Bootstrap), Vanilla JS yang akan berkomunikasi sepenuhnya lewat format JSON secara asinkron.",
            "2. Gateway & API Layer (Backend Web): Dibangun dengan Node.js (Express) atau Laravel. Memiliki tanggung jawab meng-handle HTTP request, session manajemen, logging aktivitas ke DBMS, serta memastikan keamanan file terunggah dari injeksi berbahaya (sanitize).",
            "3. AI & ML Computation Layer (AI Backend): Dibangun eksklusif menggunakan Python (Flask/FastAPI). Merupakan server spesifik yang akan memuat library Neural Network (seperti Transformers dan OpenCV) ke dalam memori RAM/GPU. Tugasnya hanya satu: Menerima file mentah, memprosesnya dengan AI cerdas, dan mereturn file matang.",
            "4. Database & Storage Layer: Penyimpanan metadata aktivitas menggunakan RDBMS (PostgreSQL/MySQL), dan penyimpanan berkas file sementara (local storage / Amazon S3 bucket)."
        ]),
        ("15. Arsitektur AI", [
            "Struktur Komponen Inti AI Engine:",
            "- Layer Ingesti Data (Parsers): Koleksi pustaka seperti PyMuPDF (membaca PDF tingkat rendah), python-docx (membaca word object tree), dan openpyxl.",
            "- Layer Pemrosesan Logika (Inference Models):",
            "  * Modul Visual (Vision Model): Menggunakan pendekatan Object Detection (seperti LayoutLMv3 atau YOLO) yang dilatih khusus guna mengidentifikasi 'Bounding Box' dari Tabel atau Gambar dalam sebuah halaman.",
            "  * Modul Tekstual (NLP): Menganalisis karakter string berhuruf tebal sebagai potensi sebuah bab (Heading), merangkai kalimat putus (Hyphenation), dan mengenali daftar isi.",
            "  * Modul Digitalisasi (OCR Engine): Untuk menterjemahkan kontur huruf dari piksel gambar ke format string.",
            "- Layer Sinkronisasi (Reconstructor Heuristics): Aturan logis (Rule-based constraints) cerdas untuk memaksa lebar tabel menyesuaikan area kertas margin agar layout tidak terlempar ke luar halaman paska transformasi."
        ]),
        ("16. DFD Level 0", [
            "Aktor Eksternal: Pengguna (User).",
            "Sistem Inti: [0.0] Sistem Smart Document Converter.",
            "Alur Data:",
            "- Pengguna memberikan input -> (Dokumen Mentah, Jenis Konversi Tujuan) menuju Proses [0.0].",
            "- Proses [0.0] memproses data dan mengirimkan output -> (Dokumen Konversi Final, Status Pratinjau/AI Summary) kepada Pengguna."
        ]),
        ("17. DFD Level 1", [
            "Rincian Proses Internal Level 1:",
            "1. Proses [1.0] Modul Penerimaan & Validasi:",
            "   - Data In: File Dokumen Mentah (User).",
            "   - Data Out: Data yang lolos validasi keamanan (dikirim ke Proses 3.0).",
            "2. Proses [2.0] Pencatatan Manajemen Basis Data:",
            "   - Menerima rincian nama file dan mencatat 'Sesi Konversi' ke D1 (Tabel Riwayat).",
            "3. Proses [3.0] Pemrosesan Cerdas AI (AI Core Processing):",
            "   - Mengambil file valid, memicu modul Layout Detection & Smart OCR, kemudian merender ulang format.",
            "   - Menghasilkan 'File Hasil'.",
            "4. Proses [4.0] Pelayanan Unduhan & Laporan:",
            "   - Mengambil file final dari [3.0], lalu memberikan respon berupa file octet-stream kepada User bersama dengan rangkuman informasi AI."
        ]),
        ("18. Context Diagram", [
            "Context Diagram identik secara struktural dengan DFD Level 0, bertindak sebagai perspektif paling atas (helicopter view).",
            "Menempatkan 'Platform Cerdas Konversi Dokumen' di tengah-tengah diagram lingkaran utama.",
            "Satu-satunya aktor yang berinteraksi langsung dari luar adalah 'User', dimana arah anak panah input memasukkan File+Instruksi, sementara anak panah output keluar membawa Dokumen Jadi+Analisis File.",
            "Entitas Internal System Storage (seperti database lokal) tersembunyi di dalam 'Black Box' ini sehingga diagram hanya fokus pada interaksi batas (boundary) antar sistem utama dan aktor."
        ]),
        ("19. Sitemap Website", [
            "Peta struktur website secara hierarki:",
            "- / (Halaman Utama / Beranda - Hook informasi fitur AI, CTA ke Converter)",
            "- /converter (Halaman Tool Inti - Area interaktif drag-and-drop unggah file)",
            "- /preview (Halaman Pratinjau Visual - Menampilkan animasi progress dan ringkasan Analisis Struktur AI. Muncul setelah proses unggah selesai)",
            "- /download (Halaman Final - Titik henti untuk mengunduh dokumen secara langsung)",
            "- /history (Halaman Riwayat - Tabel historis log aktivitas konversi untuk pengguna, berbasis cookies/session)",
            "- /about (Halaman Tentang - Latar belakang akademik, anggota kelompok, serta penjelasan teknis arsitektur AI yang digunakan pada proyek)"
        ]),
        ("20. Desain Database", [
            "Tabel `users_session`:",
            "- session_id (VARCHAR Primary Key), ip_address (VARCHAR), last_active (DATETIME)",
            "",
            "Tabel `conversion_history`:",
            "- id_konversi (BIGINT Primary Key Auto Increment)",
            "- session_id (VARCHAR Foreign Key)",
            "- original_filename (VARCHAR)",
            "- converted_filename (VARCHAR)",
            "- type_conversion (VARCHAR, Enum: pdf2docx, docx2pdf, dst.)",
            "- file_size_kb (INT)",
            "- status (VARCHAR, Enum: pending, ai_processing, success, failed)",
            "- created_at (TIMESTAMP)",
            "",
            "Tabel `ai_analysis_reports` (Fitur Tambahan Opsional):",
            "- id_report (BIGINT Primary Key)",
            "- id_konversi (BIGINT Foreign Key)",
            "- summary_text (TEXT)",
            "- object_count (JSON, misal: {tables: 2, images: 1, headers: 10})"
        ]),
        ("21. Desain API", [
            "Endpoint Utama (Base URL: /api/v1):",
            "- [POST] /upload-process",
            "  * Payload: Data tipe multipart/form-data (File object dan field `convert_type`)",
            "  * Response (Async Task / Sync Task): 200 OK dengan format JSON `{ conversion_id: 88, status: 'processing' }`",
            "",
            "- [GET] /check-status/{conversion_id}",
            "  * Fungsi: Melakukan polling/mengecek apakah proses AI telah selesai.",
            "  * Response: 200 OK dengan JSON `{ status: 'completed', ai_summary: '2 tabel dikembalikan dari PDF gambar', file_ready_url: '/api/v1/download/88' }`",
            "",
            "- [GET] /download/{conversion_id}",
            "  * Fungsi: Mendistribusikan file fisik ke klien.",
            "  * Response: Menyerahkan HTTP Header Content-Disposition attachment."
        ]),
        ("22. Wireframe Website", [
            "Konsep Tata Letak Dasar (Low-Fidelity):",
            "Halaman Utama (Home & Converter):",
            "- Bagian Header Atas: Menampilkan Logo 'SmartConverter AI' di sebelah kiri dan Navigasi Tautan (Beranda, Converter, Riwayat, Tentang) di sebelah kanan.",
            "- Bagian Hero (Headline): Teks berukuran besar 'Konversi Dokumen dengan Kecerdasan Visual AI' diposisikan rata tengah.",
            "- Area Konten Tengah (Dropzone): Sebuah kotak (card) area putus-putus besar dengan ilustrasi Awan/Dokumen, memberikan sinyal bagi pengguna untuk menjatuhkan file di area tersebut.",
            "- Bagian Kontrol Ekstra: Di bawah dropzone terdapat tombol Dropdown ganda untuk memilik format ASAL (Kiri) ke format TUJUAN (Kanan).",
            "- Bagian Footer: Copyright nama tim dan tahun.",
            "",
            "Halaman Preview AI:",
            "- Konten Kiri (Kolom Kiri): Kotak laporan 'AI Analysis'. Teks ringkasan mengenai jumlah halaman, jenis dokumen, peringatan jika terdapat tabel tersembunyi.",
            "- Konten Kanan (Kolom Kanan): Lingkaran/Bar Indikator persentase, lalu berubah menjadi ikon Ceklis hijau beserta tombol 'Download Dokumen' berukuran besar."
        ]),
        ("23. UI/UX Modern", [
            "Fokus utama UI/UX dirancang layaknya platform SaaS (Software as a Service) kelas enterprise:",
            "- Pendekatan Estetika: Mengadopsi prinsip Glassmorphism (efek transparansi kaca, blur latar belakang, outline halus) pada kartu/area dropzone.",
            "- Skema Warna (Color Palette): Background dark theme elegan (Midnight Blue/Dark Slate) dipadukan dengan aksen Gradien (Cyber Purple & Neon Blue) sebagai highlight pemrosesan Kecerdasan Buatan.",
            "- Tipografi: Menggunakan font sans-serif geometris (seperti 'Outfit', 'Inter', atau 'Roboto') untuk memberikan kesan profesional dan modern.",
            "- Micro-animations: Saat file diproses, AI memunculkan animasi pemindaian gelombang (scanning wave) atau perpindahan partikel titik secara mulus. Efek hover tombol akan menimbulkan pendar bayangan (glow shadow) merespons kursor pengguna."
        ]),
        ("24. Pembagian Modul Sistem", [
            "Proyek dipecah ke dalam tiga modul logis untuk kemudahan pengelolaan secara kolaboratif:",
            "1. Modul Frontend (Client Interface Module): Bagian eksterior yang sepenuhnya menyajikan visualisasi UI/UX, routing antar halaman web, animasi interaktif, responsivitas mobile, dan mekanisme HTTP call ke Backend.",
            "2. Modul Backend (Gateway & Database Module): Penengah dari sistem yang menangani logika I/O network, security file upload, session tracking, penyimpanan database ke tabel log, dan rute delegasi.",
            "3. Modul AI Core (Intelligent Processing Module): Modul spesialis berbentuk engine tertutup Python yang mengintegrasikan Deep Learning, Computer Vision NLP, OCR (Tesseract/PaddleOCR), Layout Classifier, dan memanipulasi byte file biner secara langsung (Pandoc/Python-docx/PyMuPDF)."
        ]),
        ("25. Pembagian Tugas 3 Anggota Kelompok", [
            "Sesuai dengan pembagian 3 modul besar di atas, tugas anggota kelompok dispesifikkan menjadi:",
            "1. Ramadhan Al Fitra (Machine Learning & AI Engineer): Fokus membangun 'Modul AI Core'. Bertanggung jawab meriset model algoritma AI (seperti menguji pipeline LayoutLM dan OCR), menyusun rule Smart Formatting, dan membangun script API Python sebagai penerima job konversi.",
            "2. Taufiq (Backend Developer & System Architect): Fokus pada 'Modul Backend'. Mendesain struktur tabel database (ERD), membuat RESTful API (menggunakan Laravel atau Node.js), mengatur alur komunikasi dari server web ke server AI, serta memastikan upload dokumen berjalan aman.",
            "3. Zikri (Frontend Developer & UI/UX Designer): Fokus menangani 'Modul Frontend'. Menerjemahkan konsep Glassmorphism ke dalam desain purwarupa UI (Figma), menyusun sintaks HTML/CSS modern, mengatur animasi interaksi web, dan mengelola integrasi state upload di browser."
        ]),
        ("26. Roadmap Pengerjaan", [
            "Estimasi timeline proyek akademik selama 14 minggu pengerjaan:",
            "- Minggu 1-2 (Inisiasi & Riset AI): Menganalisis pustaka Python yang relevan (python-docx, paddleOCR), pengumpulan sampel dokumen (PDF kompleks, file Word cacat).",
            "- Minggu 3-4 (Perancangan Sistem & Desain): Zikri menyelesaikan desain UI (Figma). Taufiq menyelesaikan desain Database, ERD, dan rancangan dasar API.",
            "- Minggu 5-8 (Implementasi Kode Inti AI): Ramadhan membangun arsitektur AI (Pipeline Layout Detection, OCR Extraction, Heuristic Formatting).",
            "- Minggu 9-10 (Pengembangan Web Interaktif): Membangun Backend HTTP Server & integrasi layout HTML/CSS Frontend oleh Zikri dan Taufiq.",
            "- Minggu 11-12 (Sistem Integrasi Berkelanjutan): Menyambungkan frontend dengan API, mengkoneksikan backend PHP/Node.js dengan Python AI Engine. Pengujian End-to-End dengan dokumen asli.",
            "- Minggu 13-14 (Stabilisasi & Dokumentasi): Bug fixing pada layout dokumen yang pecah, optimasi UX, hosting deployment terbatas, dan penyusunan naskah proposal/dokumen akhir proyek."
        ]),
        ("27. Algoritma AI yang digunakan beserta alasannya", [
            "Pendekatan Artificial Intelligence pada sistem tidak bergantung pada satu algoritma tunggal melainkan pipeline hibrida (gabungan):",
            "1. Computer Vision (YOLO - Object Detection) / LayoutLMv3: Alasan utamanya adalah pustaka parsing statis sangat buruk dalam mendeteksi kotak-kotak tabel atau tata letak kompleks dalam PDF. Dengan mengandalkan deteksi objek (CV) dan pemahaman posisi karakter (LayoutLM), AI mengetahui area spasial yang direservasi sebagai Tabel atau Gambar secara matematis.",
            "2. NLP Text Hierarchy Heuristics: Menggunakan algoritma kecerdasan tekstual (memanfaatkan Spacy atau rule-based kustom) guna mempelajari bentuk gaya teks. Jika satu kalimat pendek berhuruf ukuran 16 pt tebal yang diakhiri garis baru, maka AI mengklasifikasikannya sebagai 'Heading'.",
            "3. PaddleOCR (Optical Character Recognition): Algoritma jaringan saraf Convolutional (CNN) ini terbukti jauh lebih cepat dan akurat dalam mendeteksi teks bahasa Indonesia dari dokumen gambar yang miring atau kotor ketimbang pustaka standar konvensional (seperti Tesseract standar)."
        ]),
        ("28. Dataset yang diperlukan", [
            "Dalam melakukan proses iterasi dan verifikasi mutu AI Engine, tim membutuhkan suplai dataset berikut:",
            "- PubLayNet: Merupakan benchmark standar dataset berukuran besar untuk dokumentasi pemahaman tata letak (Document Layout Analysis), sangat cocok untuk melatih model mendeteksi bounding box heading, paragraf, dan list angka.",
            "- RVL-CDIP: Koleksi ratusan ribu pindaian/scan dari dokumen hitam putih seperti formulir pendaftaran, email fisik, laporan keuangan (mendukung pengujian OCR yang andal).",
            "- Dataset Kustom Mahasiswa: Kumpulan minimal 100 sampel real file (seperti paper, surat resmi dari institusi berbahasa Indonesia, skripsi, modul berkolom ganda) yang digunakan sebagai basis validasi Smart Formatting agar sistem AI Engine tidak bias terhadap pola penulisan akademis asing."
        ]),
        ("29. Cara Training AI (jika diperlukan)", [
            "Mayoritas kapabilitas sistem bersandar pada integrasi Model AI siap pakai (Pre-trained Models/Transfer Learning). Apabila dibutuhkan kustomisasi akurasi atau deteksi khusus elemen lokal, alur Training/Fine-Tuning akan melibatkan:",
            "1. Data Annotation: Menandai area tabel spesifik atau elemen tanda-tangan kop surat dengan platform open-source Label Studio.",
            "2. Fine-Tuning Execution: Menjalankan fine-tuning berbasis framework HuggingFace/PyTorch untuk model seperti LayoutLMv3, dijalankan dalam sesi cloud (seperti layanan Google Colab).",
            "3. Inference Optimization: Membekukan bobot (Weight freezing) menjadi format eksekusi efisien yang hemat RAM saat di-load di server Python Flask produksi.",
            "4. Threshold Calibration: Melakukan penyesuaian confidence score dari OCR, agar huruf yang di-scan buram tidak diekstrak menjadi karakter aneh."
        ]),
        ("30. Penjelasan mengapa project ini termasuk Artificial Intelligence, bukan sekadar website converter biasa", [
            "Converter biasa di internet (seperti pustaka unoconv atau php-pdf) hanya mengimplementasikan parser algoritma satu arah yang bersifat statis, kaku (rigid), berbasis aturan translasi kode XML. Converter ini akan 100% gagal jika file yang dikonversi berupa PDF gambar (hasil print yang difoto) atau jika dokumen tersebut tidak mematuhi standar formating rapi (menggunakan spasi sebagai pengganti tab).",
            "",
            "Proyek ini dengan sah diklasifikasikan sebagai Artificial Intelligence karena:",
            "1. Cognitive Semantic Vision: AI secara visual mampu 'melihat' halaman sebagai kanvas spasial, mengidentifikasi bahwa sekumpulan garis membentuk Tabel, persis cara mata dan otak manusia bekerja, bukannya sekadar membaca kode binernya.",
            "2. Penggunaan Jaringan Saraf Tiruan (Deep Learning): Menerapkan Neural Networks dan OCR secara dinamis untuk mengembalikan teks dari objek piksel mati.",
            "3. Kecerdasan Generatif untuk Rekonstruksi (Smart Formatting Agent): AI diberikan kemampuan kognitif logis untuk mengambil keputusan. Sebagai contoh, saat mentransformasi Excel besar menjadi Word ukuran A4, AI akan menganalisis lebar tabel secara otomatis; jika terdeteksi melewati margin A4, AI akan berinisiatif mengecilkan font tabel tanpa diperintah (Self-adjusting formatting). Sifat adaptabilitas inilah yang merepresentasikan wujud Kecerdasan Buatan tingkat mahir."
        ])
    ]

    for title, lines in sections:
        heading = doc.add_heading(title, level=1)
        heading.paragraph_format.space_before = Pt(14)
        heading.paragraph_format.space_after = Pt(6)
        for line in lines:
            if line.startswith("- "):
                doc.add_paragraph(line[2:], style='List Bullet')
            elif line.startswith("  * "):
                p = doc.add_paragraph(line[4:], style='List Bullet 2')
                p.paragraph_format.left_indent = Inches(0.5)
            elif line.startswith("1. ") or line.startswith("2. ") or line.startswith("3. ") or line.startswith("4. ") or line.startswith("5. ") or line.startswith("6. ") or line.startswith("7. ") or line.startswith("8. ") or line.startswith("9. ") or line.startswith("10. ") or line.startswith("11. ") or line.startswith("12. "):
                doc.add_paragraph(line, style='List Number')
            else:
                p = doc.add_paragraph(line)
                p.paragraph_format.space_after = Pt(6)
                
        doc.add_paragraph()

    out_path = r"c:\Users\THINKPAD T14\OneDrive\Documents\SEMESTER 4\AI\Project\Dokumen Konsep Proyek.docx"
    
    try:
        doc.save(out_path)
        print("Dokumen berhasil dibuat dan disimpan di:", out_path)
    except Exception as e:
        print("Error saving document:", e)
        sys.exit(1)

if __name__ == "__main__":
    main()
